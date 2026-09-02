"""Normalize job-description request bodies into plain JD text.

Each supported format is a *parser* registered in ``_PARSERS`` under its
media type; unregistered types fall back to raw UTF-8 text, so a pasted JD
(or ``curl --data-binary @jd.txt``) needs no special header. Binary formats
that arrive with a generic/unknown content type are detected by magic bytes
in :func:`_sniff_media_type`.

Supported today: JSON, PDF, DOCX, plain text, Markdown.
"""

from __future__ import annotations

from collections.abc import Callable
from dataclasses import dataclass
from io import BytesIO

import docx
from pydantic import ValidationError
from pypdf import PdfReader

from services.portfolio.schemas.tailor import TailorRequest


# Uploads larger than this are rejected — jobs, like CVs, are small text.
MAX_JD_PAYLOAD_BYTES = 10 * 1024 * 1024


class PayloadTooLargeError(ValueError):
    """Body exceeds the JD payload cap; the route maps this to 413."""


@dataclass(frozen=True)
class JdInput:
    jd_text: str
    title: str = ""


def _check_size(body: bytes) -> None:
    if len(body) > MAX_JD_PAYLOAD_BYTES:
        raise PayloadTooLargeError(
            f"Job description exceeds the {MAX_JD_PAYLOAD_BYTES // (1024 * 1024)} MB limit"
        )


def _from_json(body: bytes, title: str) -> JdInput:
    try:
        payload = TailorRequest.model_validate_json(body)
    except ValidationError:
        raise ValueError(
            "JSON payload must be an object with a jd_text field"
        ) from None
    return JdInput(payload.jd_text, payload.title or title)


def _from_pdf(body: bytes, title: str) -> JdInput:
    try:
        reader = PdfReader(BytesIO(body))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        raise ValueError(f"Could not read PDF: {exc}") from exc
    return JdInput(_normalize_text(text), title)


def _from_docx(body: bytes, title: str) -> JdInput:
    try:
        document = docx.Document(BytesIO(body))
    except Exception as exc:
        raise ValueError(f"Could not read .docx: {exc}") from exc
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return JdInput(_normalize_text("\n".join(parts)), title)


def _from_text(body: bytes, title: str) -> JdInput:
    return JdInput(body.decode("utf-8", errors="replace").strip(), title)


def _normalize_text(raw: str) -> str:
    """Collapse blank noise; keep the JD readable for the matcher."""
    clean = "\n".join(line.strip() for line in raw.splitlines() if line.strip())
    return clean.strip()


Parser = Callable[[bytes, str], JdInput]

_DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

_PARSERS: dict[str, Parser] = {
    "application/json": _from_json,
    "application/pdf": _from_pdf,
    _DOCX_MEDIA_TYPE: _from_docx,
    "text/plain": _from_text,
    "text/markdown": _from_text,
    "text/x-markdown": _from_text,
}

_GENERIC_MEDIA_TYPES = {
    "",
    "application/octet-stream",
    "application/binary",
    "*/*",
}


def _sniff_media_type(body: bytes) -> str | None:
    """Detect binary formats from magic bytes (client headers lie)."""
    if body.startswith(b"%PDF"):
        return "application/pdf"
    if body.startswith(b"PK\x03\x04"):
        return _DOCX_MEDIA_TYPE
    return None


def parse_jd_input(body: bytes, content_type: str, title: str = "") -> JdInput:
    """Turn a request body into JD text plus an optional title override.

    Explicit media types are honored first; a generic/unknown content type
    (or none) is followed by magic-byte sniffing for binary formats, and
    anything else is treated as raw text. The size cap is enforced here —
    upfront, before any format parsing.
    """
    _check_size(body)
    media_type = content_type.split(";")[0].strip().lower()
    parser = _PARSERS.get(media_type)
    if parser is None and media_type in _GENERIC_MEDIA_TYPES:
        sniffed = _sniff_media_type(body)
        if sniffed:
            parser = _PARSERS[sniffed]
    if parser is not None:
        return parser(body, title)
    return _from_text(body, title)
